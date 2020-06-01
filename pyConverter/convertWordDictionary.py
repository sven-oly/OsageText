# -*- coding: utf-8 -*-

from __future__ import absolute_import, division, print_function

import os
import re
import sys

# Read and process MS Word documents, converting Old Osage encoding into
# Unicode characters.

# https://openpyxl.readthedocs.io/en/default/tutorial.html

# TIMESTAMP for version information.
TIMESTAMP = "Version 2017-12-19 13:45"

OfficialOsageFont = 'Official Osage Language'
QuinteroDictionaryOsageFont = 'Ss Do SILDoulos Q'
FONTS_TO_CONVERT = [QuinteroDictionaryOsageFont]

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx import shared

import convertUtil
#import osageConversion
import quinteroConversion

extractedFileName = False

# Flag for handling all characters in an Old font.
convertAllInOldFontRange = True

angle_pattern = r'(<[^>]*>)'
angle_prog = prog = re.compile(angle_pattern)

# Converts Osage text that is labeld with the old font encoding
# It assumes that the font has been detected.
def checkAndConvertText(textIn):

  if textIn[0] == '=':
    # Ignore function calls
    return textIn

  # Handle text in font-labeled regions
  result = quinteroConversion.quiteroOsageToUnicode(textIn)
  return result

# Fixes contents of string in angle brackets with SIL font replacements
def fixAngleString(match):
  return quinteroConversion.convertSILDoulousQtoUnicode(match.group(1))

def convertDoc(doc, unicodeFont, debugInfo=None,
               extractedFile=None):
  sections = doc.sections
  print ('  %d sections' % len(sections))

  paragraphs = doc.paragraphs
  print ('  %d paragraphs' % len(doc.paragraphs))

  print ('  %d tables' % len(doc.tables))
  if debugInfo and doc.inline_shapes:
    print ('  %d inline_shapes' % len(doc.inline_shapes))

  if debugInfo:
#    print ('    doc dir: %s' % dir(doc))
    for section in sections:
      print ('Section = %s' % section)

  numConverts = 0
  notConverted = 0
  paraNum = 0

  # Output line for extracted text and Python conversion
  extractedLine = []
  extractedCount = 0

  missingConversions = set([])
  for para in paragraphs:

    #para_format = para.paragraph_format
    para_style = para.style
    #para_alignment = para.alignment
    #para_part = para.part

    para_style_id = para_style.style_id
    runs = para.runs
    print ('    %d runs in paragraph' % (len(runs)))
    #print ('    paragraph text = %s' % (para.text))
    runNum = 0
    runNum = 1
    afterVariant = False  # For bold words after "variant of" or "variant:"

    for run in runs:
      if len(run.text):
        thisText = run.text
        originalText = thisText
        fontObj = run.font
        fontName = fontObj.name
        if not fontName:
          fontName = para_style.font.name

        # Fix SIL font encoded text within angle brackets
        thisText = angle_prog.sub(fixAngleString, thisText)

        ## DECIDE IF further CONVERSION IS NEEDED
        convertedText = thisText
        if (((runNum == 1 and run.font.bold) or
             (afterVariant and run.font.bold) or run.font.italic) and
            fontName in FONTS_TO_CONVERT and thisText):

          # Omit conversions within angle brackets.
          # For each in the split, check if it matches
          # if not, convert, else just copy
          sp = angle_prog.split(convertedText)
          pieces = []
          convertItem = True
          for s in sp:
            if angle_prog.match(s):
              pieces.append(s)
            else:
              if s != '':
                # See it it contains a grammatical term or "do not convert"
                for term in quinteroConversion.grammatical_terms:
                  if re.search(term, s):
                    convertItem = False
                for term in quinteroConversion.do_not_convert_italics:
                  if re.search(term, s):
                    convertItem = False

                if convertItem:
                  newText, notFound = checkAndConvertText(s)
                else:
                  newText = s
                pieces.append(newText)
                if notFound:
                  print('NOT FOUND in string %s: %s' % (thisText, notFound))
                  for a in notFound:
                    missingConversions.add(a)

          convertedText = ''.join(pieces)
          afterVariant = False
          #print('!!!!!!!! CONVERTING %s to %s' % (originalText, convertedText))
        else:
          # Check for variant
          m = re.search(r'\(variant', thisText)
          if (m):
            # Remember that variant precedes the next bold item.
            afterVariant = True
          else:
            afterVariant = False  # Reset

        if originalText != convertedText:
          numConverts += 1
          run.text = convertedText
          if convertedText != thisText:
            # It was changed from ASCII to Osage
            fontObj.name = unicodeFont
          if extractedFile:
            extractedFile.write('%s\t%s\t%s\n' % (
                fontName, thisText, convertedText))
        else:
            notConverted += 1
      runNum += 1
    if para.text:
      print ('    paragraph text = %s' % (para.text))
    paraNum += 1

  print('All missing characters = %s' % missingConversions)
  if extractedFile:
     extractedFile.close()

  print ('  %d values converted to Unicode' % numConverts)
  return (numConverts, notConverted)


# Process one DOCX, substituting the
def convertOneDoc(path_to_doc, unicodeFont='Noto Sans Osage',
                  outpath=None, isString=False):

  print('TIMESTAMP: convertDoc: %s, osageConversion: %s' % (
      TIMESTAMP, quinteroConversion.TIMESTAMP))
  print ('Converting Osage in file: %s' % path_to_doc)


  doc = Document(path_to_doc)

  # Set up the Unicode font.
  styles = doc.styles
  style_obj = styles.add_style('Osage_font', WD_STYLE_TYPE.CHARACTER)
  obj_font = style_obj.font
  basic_font_size = 10
  obj_font.size = shared.Pt(basic_font_size)
  obj_font.name = unicodeFont

  if extractedFileName:
    extractedFile = codecs.open(extractedFileName, 'w', 'utf-8')
  else:
    extractedFile = None

  (numConverts, numNoteConverted) = convertDoc(doc, unicodeFont, debugInfo=True,
                                               extractedFile =extractedFile)

  if numConverts:
    newName = os.path.splitext(path_to_doc)[0]
    unicode_path_to_doc = newName + '.unicode.docx'
    doc.save(unicode_path_to_doc)
    print ('  ** Saved new version to file %s\n' % unicode_path_to_doc)
  else:
    print ('  @@@ No conversion done, so no new file created.\n')


def processArgs(argv):
  if len(sys.argv) <= 1:
    print ('Usage:')
    print ('  convertWordOsage.py inputFile.docx')
    print ('  convertWordOsage.py inputFile1.docx inputFile2.docx ... ')
    print ('  convertWordOsage.py -i fileWithFileNames')
    return None

  path_to_docs = []

  if len(argv) == 2:
    path_to_docs.append(sys.argv[1])
  else:
    if len(argv) == 3 and argv[1] == '-f':
      # Get the file containing conversion list and get all items.
      path_to_docs = convertUtil.infileToList(argv[2])
      if not path_to_docs:
        print ('Error: no contents found in file %s' %
               argv[2])
        return
    else:
      # Expect a list of files in the
      path_to_docs = [path for path in argv[1:]]

  return path_to_docs


def main(argv):

  if len(sys.argv) > 1:
    path_to_doc = sys.argv[1]
  else:
    path_to_doc = 'Harry Red Eagle Track 5.doc'

  doc_list = processArgs(argv)

  convertFileCount = 0
  for doc_path in doc_list:
    convertOneDoc(doc_path)
    convertFileCount += 1
  print ('%d files were processed' % convertFileCount)


if __name__ == "__main__":
  main(sys.argv)
